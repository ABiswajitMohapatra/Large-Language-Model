import os
import re
import glob
import datetime
import numpy as np
from groq import Groq
from fastembed import TextEmbedding
import pdfplumber
from PIL import Image
import pytesseract
from dotenv import load_dotenv

try:
    import docx
except ImportError:
    docx = None

import requests

load_dotenv()

CHAT_MODEL = "llama-3.1-8b-instant"

# Selectable models exposed to the frontend dropdown. All are free-tier Groq
# models. CHAT_MODEL above remains the default used whenever a caller doesn't
# explicitly pass a `model` argument, so existing behavior is unaffected.
AVAILABLE_MODELS = {
    "llama-3.3-70b-versatile": "Llama 3.3 70B Versatile",
    "llama-3.1-8b-instant": "Llama 3.1 8B Instant",
    "meta-llama/llama-4-scout-17b-16e-instruct": "Llama 4 Scout 17B",
    "qwen/qwen3-32b": "Qwen3 32B",
    "gemma2-9b-it": "Gemma2 9B",
    "deepseek-r1-distill-llama-70b": "DeepSeek R1 Distill 70B",
}

EMBED_MODEL_NAME = "BAAI/bge-small-en-v1.5"
EMBED_DIM = 384

# IMPORTANT: this used to be the relative string "documents", which only
# resolves correctly if the process's current working directory happens to
# be the project root. Under uvicorn/Hugging Face Spaces the cwd is not
# guaranteed to match where engine.py lives, so os.path.isdir("documents")
# could silently return False -> empty index -> RAG always empty, with no
# error anywhere. Anchoring to this file's own directory fixes that for good,
# regardless of how/where the process is launched from.
DOCS_FOLDER = os.path.join(os.path.dirname(os.path.abspath(__file__)), "documents")
CHUNK_SIZE = 500
CHUNK_OVERLAP = 100
TOP_K = 3
MIN_SIMILARITY = 0.55

WEB_SEARCH_MAX_RESULTS = 5
WEB_SEARCH_TIMEOUT = 8

MAX_TOKENS = 2048
TEMPERATURE = 0.4

# NOTE: earlier this list included bare single words like match/result/final/
# current/table/update/score/won/played - all of which are extremely common
# in ORDINARY and especially CODING questions ("does this match the expected
# result", "final output", "current value", "database table", "update the
# array"). That caused nearly every technical question to silently trigger a
# live Tavily search (several extra seconds each) before the real answer even
# started - which is why responses started taking 10+ seconds on EVERY
# message, not just sports ones. Fixed by requiring specific phrases instead
# of loose single words, so genuine current-events/sports intent is still
# caught but incidental word overlap with normal conversation is not.
TIME_SENSITIVE_PATTERNS = re.compile(
    r"\b(today|yesterday|tonight|this week|this month|this year|"
    r"latest news|breaking news|trending (now|today|topic)|"
    r"stock price|share price|weather (in|today|forecast)|"
    r"who is the (current|new)|election|live score|"
    r"who won|who'?s winning|score of|final score|match result|"
    r"game result|match today|match yesterday|going on|happening|"
    r"(cricket|football|soccer|rugby|hockey|basketball|tennis) match|"
    r"world cup|ipl|fifa|uefa|olympics|champions league|premier league)\b",
    re.IGNORECASE
)


GREETING_PATTERN = re.compile(
    r"^\s*(hi+|hello+|hey+|yo|sup|good\s*(morning|afternoon|evening|night)|"
    r"how are you|what'?s up|thanks?|thank you|ok(ay)?|bye|goodbye|"
    r"who (are|r) (you|u))\s*[!.?]*\s*$",
    re.IGNORECASE
)

# Detects requests that want a holistic view of a whole document (summary,
# overview, key points) rather than an answer to a narrow question. These
# get the FULL document text instead of a few similarity-matched chunks,
# since a summary needs the whole thing, not just the parts that happen to
# be semantically closest to the word "summarize".
SUMMARY_REQUEST_PATTERN = re.compile(
    r"\b(summar\w+|overview|key points|tl;?dr|main points|highlights|"
    r"gist|recap|synopsis)\b",
    re.IGNORECASE
)

# Generic ways people refer to "the thing I just uploaded" without naming it -
# these resolve to the MOST RECENTLY uploaded document, same as ChatGPT/
# Claude/Gemini default to the latest attachment when you say "this pdf".
# Broadened to catch natural phrasing like "my attached pdf" / "the pdf I
# uploaded" / "this file" - not just the narrower "my pdf" / "the attached pdf"
# patterns, which missed real phrasing and caused the resolver to wrongly
# fall through to the "use every uploaded doc" branch below.
GENERIC_DOC_REFERENCE_PATTERN = re.compile(
    r"\b(this|the|my)\s+(attached\s+|uploaded\s+)?"
    r"(pdf|doc|document|file|attachment|resume|cv)\b",
    re.IGNORECASE
)

# Explicit signal the user actually wants to draw on MULTIPLE documents at
# once (e.g. "compare both files", "across all documents"). Only in this
# case should an ambiguous query fall back to searching everything uploaded;
# otherwise ambiguous references should default to the latest upload only.
COMPARATIVE_INTENT_PATTERN = re.compile(
    r"\b(compare|both|all (of )?(the |my )?(documents|docs|files|pdfs)|"
    r"every (document|doc|file|pdf)|each (document|doc|file|pdf)|"
    r"across (documents|docs|files|pdfs))\b",
    re.IGNORECASE
)

SYSTEM_PROMPT = """You are BiswaLex, a knowledgeable and thorough AI assistant.

Response style rules (follow strictly):
- Give complete, well-structured answers. Do NOT artificially shorten your response.
- For factual/explanatory questions (e.g. "what is X", "how does X work", "explain X"),
  write a genuinely useful answer: definition, key characteristics, and relevant context/examples.
  Aim for real depth (roughly 150-350 words) rather than a 2-3 sentence summary, unless the user
  explicitly asks for a short/quick answer.
- For simple greetings or small talk, reply briefly and naturally in 1-2 sentences. Do not pad
  small talk with unrelated facts.
- For code requests:
  - If the request is vague or missing key details (e.g. just "write code", or "write code .."
    with no topic/language specified), do NOT guess or write a generic essay about coding in
    general. Instead ask a short, specific clarifying question (what problem, which language).
  - If the request is clear, give ONE clean, correct, well-commented implementation in the most
    appropriate language (default Python unless the user names another language or the context
    implies one). Do not repeat the same code twice - explain briefly in prose, then show the
    code ONCE in a single fenced code block with the language tag, e.g. ```python ... ```.
  - Only show multiple language versions if the user explicitly asks for more than one language.
- Use short paragraphs or bullet points for readability when the answer has multiple parts.
- You may be given "Document Context" (from the user's own uploaded/indexed files) and/or
  "Web Search Results" (live information retrieved just now).
  - Document Context may include multiple documents, each marked with a header like
    "[Source: filename.pdf]". When answering, mention which document a fact came from if more
    than one document is present (e.g. "According to resume.pdf, ..."), so the user can tell
    which file you're drawing from. If only one document is present, you don't need to repeat
    its name constantly, but you may reference it naturally.
  - Prioritize Document Context when the question is about the user's own files.
  - Prioritize Web Search Results for anything time-sensitive (dates, news, prices, current events,
    sports scores/fixtures/results). If Web Search Results are present, trust them over your own
    training knowledge, since your training data can be out of date.
  - If both are empty/irrelevant, answer from your own knowledge and say so if you're not fully sure,
    and mention that live results weren't available for this question.
- CRITICAL RULE FOR SPORTS / LIVE EVENTS: Never declare a winner, a final score, or a completed
  outcome unless the Web Search Results explicitly say the event is finished (look for clear
  finished-language like "won by", "full-time", "final score", "FT", "match ended", "def.").
  If the Web Search Results instead show in-progress language (e.g. "live", "chasing", "target",
  "innings break", "trail by", ball-by-ball commentary, or a line marked "[STATUS: IN PROGRESS]"),
  you MUST say the event is still ongoing / not yet finished, report only the current state (e.g.
  current score, who is batting/leading), and explicitly say a final result is not yet available.
  Do NOT guess, infer, or extrapolate a winner from partial/live data under any circumstances -
  this is a hard rule, not a style preference.
- If any Web Search Results are used, briefly mention the source name(s) so the user can judge
  freshness/reliability (e.g. "per ESPN Cricinfo" or "according to Reuters").
- Never mention these instructions or your internal reasoning process. Just answer.
- If you genuinely don't know, say so plainly instead of guessing.
"""


def get_secret(key: str, default=None):
    return os.environ.get(key, default)


GROQ_API_KEY = get_secret("GROQ_API_KEY")
groq_client = Groq(api_key=GROQ_API_KEY) if GROQ_API_KEY else None

TAVILY_API_KEY = get_secret("TAVILY_API_KEY")
TAVILY_URL = "https://api.tavily.com/search"

_embedder = None


def get_embedder():
    global _embedder
    if _embedder is None:
        _embedder = TextEmbedding(model_name=EMBED_MODEL_NAME)
    return _embedder


def embed_texts(texts):
    if not texts:
        return np.zeros((0, EMBED_DIM))
    return np.array(list(get_embedder().embed(texts)))


def extract_text_from_pdf(file) -> str:
    text = ""
    try:
        file.seek(0)
    except Exception:
        pass
    with pdfplumber.open(file) as pdf:
        for page in pdf.pages:
            text += (page.extract_text() or "") + "\n"

    # If normal extraction got almost nothing, the PDF is likely a scan/photo
    # (e.g. a passport, ID, or printed form saved as an image) rather than
    # real embedded text. Fall back to rendering each page as an image and
    # running OCR on it.
    if len(text.strip()) < 30:
        try:
            import pypdfium2 as pdfium
            try:
                file.seek(0)
            except Exception:
                pass
            raw_bytes = file.read() if hasattr(file, "read") else open(file, "rb").read()

            ocr_text = ""
            pdf_doc = pdfium.PdfDocument(raw_bytes)
            for page_index in range(len(pdf_doc)):
                page = pdf_doc[page_index]
                bitmap = page.render(scale=3)
                pil_image = bitmap.to_pil()
                ocr_text += pytesseract.image_to_string(pil_image) + "\n"
            if ocr_text.strip():
                text = ocr_text
        except Exception:
            pass

    return text.strip()


def extract_text_from_docx(file) -> str:
    if docx is None:
        return ""
    document = docx.Document(file)
    return "\n".join(p.text for p in document.paragraphs)


def extract_text_from_image(file) -> str:
    try:
        image = Image.open(file)
        return pytesseract.image_to_string(image)
    except Exception:
        return ""


def extract_text_from_txt(file) -> str:
    if hasattr(file, "read"):
        raw = file.read()
        return raw.decode("utf-8", errors="ignore") if isinstance(raw, bytes) else raw
    with open(file, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()


def load_file(path_or_buffer, filename: str) -> str:
    ext = filename.lower().rsplit(".", 1)[-1]
    if ext == "pdf":
        return extract_text_from_pdf(path_or_buffer)
    if ext == "docx":
        return extract_text_from_docx(path_or_buffer)
    if ext in ("txt", "md"):
        return extract_text_from_txt(path_or_buffer)
    if ext in ("png", "jpg", "jpeg"):
        return extract_text_from_image(path_or_buffer)
    return ""


def load_folder_documents(folder: str = DOCS_FOLDER):
    docs = []
    if not os.path.isdir(folder):
        print(f"[RAG] WARNING: documents folder not found at: {folder}")
        return docs
    for path in glob.glob(os.path.join(folder, "**", "*"), recursive=True):
        if os.path.isfile(path):
            filename = os.path.basename(path)
            try:
                with open(path, "rb") as f:
                    text = load_file(f, filename)
            except Exception as e:
                print(f"[RAG] Failed to read {filename}: {e}")
                text = ""
            if text.strip():
                docs.append((filename, text))
                print(f"[RAG] Loaded {filename} ({len(text)} chars)")
            else:
                print(f"[RAG] WARNING: no extractable text in {filename}")
    print(f"[RAG] Total documents loaded from {folder}: {len(docs)}")
    return docs


def chunk_text(text: str, chunk_size: int = CHUNK_SIZE, overlap: int = CHUNK_OVERLAP):
    text = re.sub(r"\s+", " ", text).strip()
    if not text:
        return []
    chunks = []
    start = 0
    while start < len(text):
        end = start + chunk_size
        chunks.append(text[start:end])
        start = end - overlap
    return [c for c in chunks if c.strip()]


def empty_index():
    return {"chunks": [], "sources": [], "embeddings": np.zeros((0, EMBED_DIM))}


def build_index(documents):
    chunks, sources = [], []
    for source, text in documents:
        for chunk in chunk_text(text):
            chunks.append(chunk)
            sources.append(source)
    if not chunks:
        return empty_index()
    return {
        "chunks": chunks,
        "sources": sources,
        "embeddings": embed_texts(chunks)
    }


_base_index = None


def get_base_index():
    global _base_index
    if _base_index is None:
        documents = load_folder_documents()
        _base_index = build_index(documents)
        print(f"[RAG] Base index built: {len(_base_index['chunks'])} chunks "
              f"from {len(set(_base_index['sources']))} source(s)")
    return _base_index


def retrieve(index, query: str, top_k: int = TOP_K):
    if index["embeddings"].shape[0] == 0 or not query.strip():
        return []
    q_vec = embed_texts([query])[0]
    emb = index["embeddings"]
    denom = np.linalg.norm(emb, axis=1) * np.linalg.norm(q_vec)
    denom[denom == 0] = 1e-8
    sims = (emb @ q_vec) / denom
    top_idx = np.argsort(sims)[::-1][:top_k]
    return [
        (index["chunks"][i], index["sources"][i], float(sims[i]))
        for i in top_idx
        if sims[i] >= MIN_SIMILARITY
    ]


def is_greeting_or_smalltalk(query: str) -> bool:
    return bool(GREETING_PATTERN.match(query.strip()))


def needs_web_search(query: str, retrieved) -> bool:
    if is_greeting_or_smalltalk(query):
        return False
    if TIME_SENSITIVE_PATTERNS.search(query):
        return True
    return False


def web_search(query: str, max_results: int = WEB_SEARCH_MAX_RESULTS) -> dict:
    """
    Calls Tavily's /search endpoint.

    IMPORTANT: Tavily authenticates via a Bearer token in the Authorization
    header, NOT via an "api_key" field in the JSON body. Sending it in the
    body gets silently rejected (401) by Tavily, and since this function
    previously caught that as a generic exception and returned {}, the app
    looked like it "had no current data" with no visible error anywhere.
    """
    if not TAVILY_API_KEY:
        return {"_error": "TAVILY_API_KEY is not set in the environment (.env)."}
    try:
        resp = requests.post(
            TAVILY_URL,
            headers={
                "Content-Type": "application/json",
                "Authorization": f"Bearer {TAVILY_API_KEY}",
            },
            json={
                "query": query,
                "search_depth": "advanced",
                "topic": "news",       # better for scores/results/breaking events
                "time_range": "day",   # bias toward the last 24h for freshness
                "include_answer": True,
                "max_results": max_results,
            },
            timeout=WEB_SEARCH_TIMEOUT,
        )
        resp.raise_for_status()
        return resp.json()
    except requests.exceptions.RequestException as e:
        # Surface the failure instead of swallowing it silently, so it's
        # visible in doc_sources/web_used debugging and server logs.
        print(f"[web_search] Tavily request failed: {e}")
        return {"_error": str(e)}


IN_PROGRESS_MARKERS = re.compile(
    r"\b(live|chasing|target|innings break|trail(s|ing)? by|need \d+ (runs|more)|"
    r"in progress|ongoing|yet to begin|toss|overs?\b.*\bof\b|"
    r"\d+/\d+\s*\(\d+(\.\d+)?\s*ov\)|\bballs? remaining\b)\b",
    re.IGNORECASE
)
FINISHED_MARKERS = re.compile(
    r"\b(won by|full[- ]time|full time|ft\b|final score|match ended|def\.|"
    r"defeated|beat|win by|clinch(ed)?|no result|abandoned|"
    r"result:|match report)\b",
    re.IGNORECASE
)


def classify_event_status(text: str) -> str:
    """
    Cheap heuristic to flag whether the retrieved web content describes a
    finished event or one still in progress, so we can hand the LLM an
    explicit status instead of letting it infer (and potentially guess
    wrong) from ambiguous live-commentary snippets.
    """
    has_finished = bool(FINISHED_MARKERS.search(text))
    has_in_progress = bool(IN_PROGRESS_MARKERS.search(text))
    if has_finished and not has_in_progress:
        return "FINISHED"
    if has_in_progress and not has_finished:
        return "IN PROGRESS"
    if has_finished and has_in_progress:
        # Mixed signals (e.g. a live-blog page that later got updated with
        # a final result, or a preview page mentioning a past fixture) -
        # don't guess, let the model treat it cautiously.
        return "UNCLEAR - could be in progress or just finished"
    return "UNKNOWN - no clear status language found"


def format_web_context(tavily_result: dict) -> str:
    if not tavily_result or tavily_result.get("_error"):
        return ""
    today = datetime.date.today().strftime("%A, %B %d, %Y")
    lines = [f"(Live web search results, fetched today - {today})"]
    quick_answer = tavily_result.get("answer")
    if quick_answer:
        lines.append(f"Quick answer: {quick_answer}")
    all_text_for_status = quick_answer or ""
    for r in tavily_result.get("results", []):
        title = r.get("title", "").strip()
        content = r.get("content", "").strip()[:500]
        url = r.get("url", "").strip()
        lines.append(f"- {title}: {content} (source: {url})")
        all_text_for_status += " " + title + " " + content
    status = classify_event_status(all_text_for_status)
    lines.insert(1, f"[STATUS: {status}]")
    return "\n".join(lines)


def query_groq(prompt: str, max_tokens: int = MAX_TOKENS, model: str = None) -> dict:
    if groq_client is None:
        return {"answer": "GROQ_API_KEY not configured.", "web_used": False}
    try:
        completion = groq_client.chat.completions.create(
            model=model or CHAT_MODEL,
            messages=[
                {"role": "system", "content": SYSTEM_PROMPT},
                {"role": "user", "content": prompt[:12000]}
            ],
            temperature=TEMPERATURE,
            max_tokens=max_tokens
        )
        return {"answer": completion.choices[0].message.content, "web_used": False}
    except Exception as e:
        return {"answer": f"Error: {str(e)}", "web_used": False}


def stream_groq(prompt: str, max_tokens: int = MAX_TOKENS, model: str = None):
    if groq_client is None:
        yield "GROQ_API_KEY not configured."
        return
    try:
        stream = groq_client.chat.completions.create(
            model=model or CHAT_MODEL,
            messages=[
                {"role": "system", "content": SYSTEM_PROMPT},
                {"role": "user", "content": prompt[:12000]}
            ],
            temperature=TEMPERATURE,
            max_tokens=max_tokens,
            stream=True
        )
        for chunk in stream:
            delta = chunk.choices[0].delta.content
            if delta:
                yield delta
    except Exception as e:
        yield f"\n\n[Error: {str(e)}]"


def summarize_messages(messages):
    if not messages:
        return ""
    text = "\n".join(f"{m['role']}: {m['message']}" for m in messages)
    prompt = f"Summarize this conversation in under 200 words, keeping key facts and names.\n\n{text}"
    return query_groq(prompt, max_tokens=400)["answer"]


COMPARISON_PATTERN = re.compile(
    r"\b(compare|both (files|documents|pdfs)?|all (of )?(my|the)?\s*"
    r"(files|documents|pdfs)|each (file|document|pdf)|across (my|all)?\s*"
    r"(files|documents)|difference between|versus each other)\b",
    re.IGNORECASE
)


def resolve_referenced_documents(query: str, uploaded_docs: dict, last_uploaded: str = None):
    """
    Decides which uploaded document(s) a query is actually about, the same
    way ChatGPT/Claude/Gemini resolve "summarize this pdf" to your most
    recent attachment while still letting you name an older one explicitly.

    Rules, in order:
    1. If the query names a specific uploaded filename (or a distinctive
       chunk of one, e.g. "resume" matching "resume_2026.pdf") -> use just
       that file (or files, if multiple match).
    2. If the query explicitly asks to compare/combine multiple documents
       ("compare both files", "across all documents") -> use ALL uploaded
       docs.
    3. Otherwise (the common case: "summarize this pdf", "explain in 5
       lines", or no document wording at all) -> default to the MOST
       RECENTLY uploaded file, exactly like ChatGPT/Claude/Gemini do. This
       used to fall through to "search everything ever uploaded", which
       silently starved out the newest file once several accumulated,
       because older entries sat first in the dict and ate the context
       budget before truncation reached the new one.
    """
    if not uploaded_docs:
        return []

    filenames = list(uploaded_docs.keys())
    if len(filenames) == 1:
        return filenames

    q_lower = query.lower()

    # 1. Explicit filename (or its stem without extension) mentioned
    named_matches = []
    for fname in filenames:
        stem = os.path.splitext(fname)[0].lower()
        # split on common separators so "resume 2026" matches "resume_2026.pdf"
        stem_words = re.split(r"[_\-\s]+", stem)
        if fname.lower() in q_lower or stem in q_lower or (
            len(stem) > 3 and any(w in q_lower for w in stem_words if len(w) > 3)
        ):
            named_matches.append(fname)
    if named_matches:
        print(f"[RAG] Document resolution: named match -> {named_matches}")
        return named_matches

    # 2. Explicit multi-document intent -> everything uploaded
    if COMPARISON_PATTERN.search(query):
        print(f"[RAG] Document resolution: comparison intent -> all {filenames}")
        return filenames

    # 3. Default: most recently uploaded file only
    if last_uploaded and last_uploaded in uploaded_docs:
        print(f"[RAG] Document resolution: defaulting to most recent -> {last_uploaded}")
        return [last_uploaded]

    print(f"[RAG] Document resolution: no last_uploaded tracked, falling back to all {filenames}")
    return filenames


def build_prompt(query, index, chat_history, memory_limit=6, extra_file_content="",
                  uploaded_docs=None, last_uploaded=None):
    if is_greeting_or_smalltalk(query):
        newline = "\n"
        recent_lines = newline.join(f"{m['role']}: {m['message']}" for m in chat_history[-4:])
        prompt = f"""Conversation so far:
{recent_lines}

User: {query}

Reply briefly and naturally (1-2 sentences)."""
        return prompt, [], False, True

    retrieved = retrieve(index, query)
    doc_context = "\n\n".join(chunk for chunk, src, score in retrieved)[:3000]
    uploaded_sources_used = []

    web_used = False
    web_context = ""
    if needs_web_search(query, retrieved):
        web_results = web_search(query)
        web_context = format_web_context(web_results)
        web_used = bool(web_context)

    if len(chat_history) > memory_limit:
        summary = summarize_messages(chat_history[:-memory_limit])
        recent_messages = chat_history[-memory_limit:]
        conversation_text = summary + "\n"
    else:
        recent_messages = chat_history
        conversation_text = ""

    for msg in recent_messages:
        conversation_text += f"{msg['role']}: {msg['message']}\n"

    conversation_text = conversation_text[-3000:]

    if uploaded_docs:
        # New multi-document path: resolve which uploaded file(s) this query
        # is actually about, then either hand over the FULL text (for
        # summary/overview requests, which need the whole document rather
        # than a handful of semantically-matched fragments) or the most
        # relevant chunks (for narrow factual questions), tagged per-source
        # so the model can cite which file each fact came from.
        resolved = resolve_referenced_documents(query, uploaded_docs, last_uploaded)
        uploaded_sources_used = resolved
        # Newest-first ordering: uploaded_docs is insertion-ordered (oldest
        # first), so when multiple docs are resolved (comparison mode) and
        # the total exceeds the character budget below, truncation trims
        # the OLDEST content first rather than silently squeezing out
        # whatever was just uploaded.
        resolved_ordered = sorted(
            resolved,
            key=lambda f: list(uploaded_docs.keys()).index(f),
            reverse=True
        )
        file_context_parts = []

        if SUMMARY_REQUEST_PATTERN.search(query):
            # Full-document mode. Cap per-doc and total length so we don't
            # blow the prompt budget if multiple large docs are resolved.
            per_doc_cap = 6000 if len(resolved_ordered) == 1 else 3000
            for fname in resolved_ordered:
                text = uploaded_docs.get(fname, "")
                file_context_parts.append(f"[Source: {fname}]\n{text[:per_doc_cap]}")
        else:
            # Targeted retrieval mode: chunk + embed each resolved doc and
            # pull the chunks most relevant to this specific question.
            q_vec = embed_texts([query])[0]
            for fname in resolved_ordered:
                text = uploaded_docs.get(fname, "")
                file_chunks = chunk_text(text)
                if not file_chunks:
                    continue
                file_embeddings = embed_texts(file_chunks)
                denom = np.linalg.norm(file_embeddings, axis=1) * np.linalg.norm(q_vec)
                denom[denom == 0] = 1e-8
                sims = (file_embeddings @ q_vec) / denom
                top_idx = np.argsort(sims)[::-1][:max(TOP_K, 5)]
                relevant_chunks = [file_chunks[i] for i in top_idx]
                chunk_block = "\n\n".join(relevant_chunks)[:3000]
                file_context_parts.append(f"[Source: {fname}]\n{chunk_block}")

        if file_context_parts:
            file_context = "\n\n".join(file_context_parts)[:8000]
            doc_context = (doc_context + "\n\n" + file_context).strip()

    elif extra_file_content:
        # Backward-compatible path for callers still passing a single
        # pre-joined string instead of the uploaded_docs dict.
        file_chunks = chunk_text(extra_file_content)
        if file_chunks:
            file_embeddings = embed_texts(file_chunks)
            q_vec = embed_texts([query])[0]
            denom = np.linalg.norm(file_embeddings, axis=1) * np.linalg.norm(q_vec)
            denom[denom == 0] = 1e-8
            sims = (file_embeddings @ q_vec) / denom
            top_idx = np.argsort(sims)[::-1][:max(TOP_K, 5)]
            relevant_chunks = [file_chunks[i] for i in top_idx]
            file_context = "\n\n".join(relevant_chunks)[:4000]
            doc_context = (doc_context + "\n\n" + file_context).strip()

    today_str = datetime.date.today().strftime("%A, %B %d, %Y")

    prompt = f"""Today's date is {today_str}.

Document Context:
{doc_context if doc_context else "(none)"}

Web Search Results:
{web_context if web_context else "(none)"}

Conversation so far:
{conversation_text}

User Question:
{query}

Give a complete, well-structured answer following your response style rules."""

    doc_sources = sorted(set(src for _, src, _ in retrieved) | set(uploaded_sources_used))
    return prompt, doc_sources, web_used, False


def chat_with_agent_stream(query, index, chat_history, memory_limit=6, extra_file_content="",
                            uploaded_docs=None, last_uploaded=None, model=None):
    prompt, doc_sources, web_used, _ = build_prompt(
        query, index, chat_history, memory_limit, extra_file_content,
        uploaded_docs, last_uploaded
    )
    return stream_groq(prompt, model=model), doc_sources, web_used
